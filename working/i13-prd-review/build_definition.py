from copy import deepcopy
import hashlib
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile, ZipInfo

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO_ROOT = Path(__file__).resolve().parents[2]
SOURCE = REPO_ROOT / "docs" / "product" / "MBPRD-P2-I13_Video_Face_Speech_Voice_Learning_v0.2.docx"
OUTPUT = REPO_ROOT / "docs" / "product" / "MBBS-P2_INCREMENT_13_DEFINITION_DRAFT_v0.2.docx"
EXPECTED_SOURCE_SHA256 = "72413EAC39018653CA0D979B9EE82CF92AD4846508FD9B82813F17305D6F238F"


def clear_body(doc):
    body = doc._element.body
    sect = body.sectPr
    for child in list(body):
        if child is not sect:
            body.remove(child)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(twips))


def style_table(table, widths=None):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")
    if widths:
        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                set_cell_width(cell, widths[idx])
    for idx, cell in enumerate(table.rows[0].cells):
        set_cell_shading(cell, "D9EAF7")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(20, 52, 78)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        tr_pr.append(OxmlElement("w:cantSplit"))
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, text in enumerate(headers):
        table.rows[0].cells[i].text = text
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
    style_table(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def normalize_docx_container(path):
    """Make regenerated DOCX bytes deterministic by normalizing ZIP metadata."""
    temp_path = path.with_name(path.name + ".normalized")
    with ZipFile(path, "r") as source_zip, ZipFile(temp_path, "w", ZIP_DEFLATED) as target_zip:
        for source_info in sorted(source_zip.infolist(), key=lambda item: item.filename):
            target_info = ZipInfo(source_info.filename, (1980, 1, 1, 0, 0, 0))
            target_info.compress_type = ZIP_DEFLATED
            target_info.external_attr = source_info.external_attr
            target_info.create_system = source_info.create_system
            target_zip.writestr(target_info, source_zip.read(source_info.filename))
    temp_path.replace(path)


actual_source_sha256 = hashlib.sha256(SOURCE.read_bytes()).hexdigest().upper()
if actual_source_sha256 != EXPECTED_SOURCE_SHA256:
    raise RuntimeError(f"PRD source hash mismatch: {actual_source_sha256}")

doc = Document(SOURCE)
clear_body(doc)

title = doc.add_paragraph(style="Title")
title.add_run("P2-I13: Video, Face, Speech & Voice Learning")
subtitle = doc.add_paragraph(style="Subtitle")
subtitle.add_run("Increment Definition — Draft v0.2 for Founder Review")

add_table(
    doc,
    ["Document", "Value"],
    [
        ("Identifier", "MBBS-P2_INCREMENT_13_DEFINITION"),
        ("Status", "DRAFT FOR FOUNDER REVIEW — BUILD NOT AUTHORIZED"),
        ("Definition version", "Draft v0.2"),
        ("Controlling PRD", "MBPRD-P2-I13 Video Face, Speech & Voice Learning v0.2"),
        ("PRD verification", "Footer states Draft v0.2; five embedded approved-direction screens verified"),
        ("Increment", "P2-I13"),
        ("Date", "4 September 2026"),
        ("Planning boundary", "Definition review only; no implementation, migration, processing, cleanup, or archive-wide recognition"),
    ],
    [2550, 6810],
)

p = doc.add_paragraph()
r = p.add_run("Review gate. ")
r.bold = True
p.add_run("This definition translates PRD v0.2 into an increment boundary for founder review. It does not authorize a build. Explicit founder authorization is required after the definition and its open decisions are accepted.")

doc.add_heading("1. Increment outcome", level=1)
doc.add_paragraph(
    "P2-I13 revalidates MemoryBox's video, face, speech, and voice pipeline so that immutable source video remains the evidence anchor while time-addressed observations, learned exemplars, suggestions, corrections, and playback moments remain derived, auditable, and rebuildable. The complete lifecycle must be proven on a versioned bounded corpus before archive-wide recognition can be unlocked."
)
add_bullets(doc, [
    "Repair the approximately 155,000-recognition-event explosion and the 1–2 second pseudo-video fragment experience without altering source media.",
    "Open one source video at an evidence-backed time and continue natural playback beyond the matched interval.",
    "Synchronize transcript navigation and playback in both directions while preserving the original machine transcript.",
    "Support owner-confirmed face and voice learning as independent evidence that may corroborate, but never silently substitute for, one another.",
    "Make correction, retirement, downstream invalidation, and bounded reprocessing reversible and provenance-preserving.",
    "Keep archive-wide recognition locked until bounded-corpus acceptance and a separate founder unlock decision.",
    "Prove exact-phrase, Person-speaking, semantic-subject, and Person-appearance retrieval, with every result source-linked and time-addressable.",
])

doc.add_heading("2. Definition status and authority", level=1)
add_table(doc, ["Authority", "Treatment"], [
    ("MBPRD-P2-I13 v0.2", "Controlling product requirements and approved-direction screen packet."),
    ("Roadmap sequencing", "This definition does not resequence the roadmap. Current founder direction identifies I14 as Comms + Gallery; any roadmap text that differs must be reconciled separately before later-increment planning."),
    ("Existing repository behavior", "Must be assessed and classified before change; working behavior is preserved rather than rebuilt by assumption."),
    ("I13 definition", "Once founder-accepted, fixes scope, gates, exclusions, and build-authorization boundary."),
    ("Founder acceptance", "Required separately for definition, build authorization, bounded-corpus acceptance, archive unlock, and archive processing start."),
], [2700, 6660])

doc.add_heading("3. Mandatory assessment phase", level=1)
doc.add_paragraph("The first authorized implementation phase, if later approved, is assessment—not remediation or processing. It must inventory the current pipeline and classify every requirement as Complete, Partial, Defective, Missing, or Unknown.")
add_bullets(doc, [
    "Trace source-video ingestion, stable identity, provider provenance, and playback location.",
    "Trace face detection, tracking, recognition, appearance grouping, Gallery/Person consumption, and correction paths.",
    "Trace audio extraction, speech regions, STT, transcript spans, diarization, speaker teaching, voice exemplars, and retrieval.",
    "Inventory all recognition tables, indexes, queues, derivative files, APIs, jobs, and UI consumers.",
    "Reproduce the 155,000-event condition and identify whether short fragments are database records, physical derivative files, or both.",
    "Locate an existing versioned 22-video bounded-corpus manifest or propose one for separate review before any recognition run.",
])

doc.add_heading("4. Architectural invariants", level=1)
add_bullets(doc, [
    "One immutable source video may have many derived observations and moments; derived records are not new videos.",
    "No I13 process writes corrections or identity conclusions back to Immich.",
    "Original source video, source audio, and machine transcript remain preserved. Owner transcript edits are additive correction overlays.",
    "Every derived record resolves to a stable source identifier and valid time range or playhead position.",
    "Identical pipeline version, source, model, and parameter reruns are idempotent and cannot multiply equivalent records.",
    "A Person may appear without speaking and may speak off-camera. Face co-occurrence alone never proves speaker identity.",
    "Learned evidence records Person, source, interval/frame, method, quality, actor, status, configuration, and provenance.",
    "Correction supersedes prior associations without erasing history. Retirement stops future exemplar use without deleting source evidence.",
])

doc.add_heading("5. In-scope workstreams", level=1)
add_table(doc, ["Order", "Workstream", "Required result"], [
    ("1", "Defect diagnosis", "Causal analysis, backup/recovery plan, derived-data classification, reconciled counts."),
    ("2", "Safe remediation", "Quarantine or supersede invalid derived data; prevent recurrence; prove idempotency before any physical derivative deletion."),
    ("3", "Source playback", "Open the immutable source at the evidence start; continue normal playback; preserve return context."),
    ("4", "Transcript synchronization", "Playback follows transcript and transcript selection seeks video without fighting deliberate scrolling."),
    ("5", "Face pipeline", "Revalidate detection, exemplars, continuity, grouping, owner Learn, review, correction, and retirement."),
    ("6", "Voice pipeline", "Select/refine speech, preserve transcript overlay, assign Person, confirm quality, recognize off-camera speech."),
    ("7", "Corroboration", "Combine independent face/voice evidence transparently without collapsing modality boundaries."),
    ("8", "Owner administration", "Owner-only Admin navigation, Processing Jobs, Learned Evidence, Archive Health, Historian Campaigns, and setup destinations."),
    ("9", "Bounded proof", "Run only manifest-controlled sources; publish accuracy, integrity, timing, failure, and residual-risk evidence."),
    ("10", "Release control", "Reject archive scope until founder acceptance creates an unlock record; starting archive processing remains separate."),
], [700, 2300, 6360])

doc.add_heading("6. User-experience definition", level=1)
doc.add_paragraph("The five screens embedded in PRD v0.2 are the approved behavioral direction. They do not authorize a replacement design system. Implementation must reuse the MemoryBox dark shell, shared navigation, evidence viewer, Person patterns, and existing components wherever suitable.")
doc.add_paragraph("Historian Campaigns may be linked as an owner-only Admin destination, but that navigation placement does not reopen, redesign, replace, or alter the accepted P2-I12 Historian Capture workflow.")
add_bullets(doc, [
    "Video Detail / Learn supports transcript selection, face sample, voice confirmation, and a combined action while retaining independent evidence records.",
    "Video Detail / People distinguishes appearances, speech, confirmed learning, and unknown observations, all linked to source time.",
    "Admin is an owner-only top-navigation destination, not a developer console.",
    "Processing Jobs exposes bounded scope, persisted progress/counts, failures, Pause/Retry/View actions, and the enforced archive lock.",
    "Learned Evidence exposes provenance, source-time navigation, Person correction, transcript overlay editing, and reasoned reversible retirement.",
    "Before confirmation, copy must say ‘Assigned for learning: [Person]’ rather than implying recognition already proved identity.",
])

doc.add_heading("7. Bounded corpus and safety gate", level=1)
add_bullets(doc, [
    "Corpus membership is an explicit versioned manifest of stable video identifiers; no wildcard or currently reachable-file inference is allowed.",
    "Coverage includes face-only, off-camera voice, simultaneous modalities, multiple people, poor audio, occlusion, short and sustained appearances, and no-match cases.",
    "Expected identities and approximate intervals are owner-confirmed sufficiently to score results.",
    "All I13 learning and recognition jobs inherit bounded scope until the archive gate is accepted and unlocked.",
    "Passing automated tests does not unlock the archive. Founder acceptance is required to unlock; a separate deliberate action is required to start processing.",
])

doc.add_heading("8. Acceptance gates", level=1)
add_table(doc, ["Gate", "Pass condition"], [
    ("Integrity", "Repeated bounded runs do not increase equivalent observations; every derivative resolves to one source; no pseudo-videos remain in presentation."),
    ("Playback", "Sampled results open the correct source/time within agreed tolerance and continue beyond the relevance interval."),
    ("Transcript", "Follow, seek, manual scroll, selection, and correction overlay work while original STT remains preserved."),
    ("Face", "Sustained appearances group sensibly; evidence can be learned, reviewed, corrected, retired, and traced."),
    ("Voice", "Clean confirmed samples produce reviewable suggestions including off-camera speech; face is not required."),
    ("Corroboration", "Each modality remains visible and attributable when confidence is combined."),
    ("Retrieval", "Exact-phrase, Person-speaking, semantic-subject, and Person-appearance queries return relevant results; every result is linked to its immutable source and opens at a valid time address."),
    ("Legacy fragments", "Every inventoried legacy fragment is classified as migrated to a source-linked moment, quarantined with a recorded reason, or verified as a generated derivative eligible for controlled deletion. No fragment remains unclassified."),
    ("Jobs", "Displayed scope, progress, counts, failures, and actions reconcile with persisted work."),
    ("Regression", "Gallery, query context, modal navigation, Person filters, timeline range/playhead, and increasing precision remain operational."),
    ("Safety", "Pre-acceptance archive starts are rejected by API/worker; no source/provider media is mutated."),
    ("Proof", "Before/after counts, tests, screenshots, timing, failures, rollback notes, and residual risks are documented."),
], [1900, 7460])

doc.add_heading("9. Explicitly out of scope", level=1)
add_bullets(doc, [
    "Replacing Immich or writing corrections back to it.",
    "General family multi-user administration.",
    "I14 unified communications aggregation, Person-wide evidence integration, or Save as Story work.",
    "Later setting, activity, artifact, or external-historical-context learning.",
    "Automatic archive unlock or automatic processing after a passing bounded run.",
    "A full developer operations console or raw infrastructure logs in the owner UI.",
    "Reopening or redesigning the accepted P2-I12 Historian Capture workflow when Historian Campaigns is linked under Admin.",
    "Resequencing I14 or later roadmap increments through this definition.",
    "Any archive-wide cleanup, recognition, migration, or deletion under the authority of this draft definition.",
])

doc.add_heading("10. Founder decisions required before build authorization", level=1)
add_table(doc, ["Decision", "Recommended default from PRD v0.2", "Status"], [
    ("Transcript correction model", "Preserve immutable machine transcript and store owner edits as overlays for display/search/learning.", "Open for founder confirmation"),
    ("Retirement cascade", "Stop future exemplar use, retain history, mark dependent suggestions stale, and reprocess only affected bounded scope.", "Open for founder confirmation"),
    ("Bounded manifest", "Use the existing versioned 22-video manifest if valid; otherwise authorize a manifest as the first reviewed I13 artifact.", "Existence must be established"),
    ("Archive unlock", "Keep unlock as an explicit founder action after I13 acceptance; keep processing start as a second action.", "Open for founder confirmation"),
    ("Thresholds", "Make face, voice, grouping, and corroboration thresholds configurable/versioned; set policy only after bounded measurements.", "Blocks acceptance policy, not assessment"),
], [2300, 5260, 1800])

doc.add_heading("11. Authorization sequence", level=1)
for text in [
    "Founder reviews and accepts or revises this increment definition.",
    "Founder resolves the decisions in Section 10 or explicitly defers a decision to a named pre-build gate.",
    "Founder separately authorizes the assessment phase. Assessment produces the completeness matrix, manifest status, causal analysis, and implementation plan; it does not run recognition or remediate data.",
    "Founder reviews assessment evidence and separately authorizes bounded implementation and testing.",
    "Founder accepts the bounded-corpus result before archive-wide recognition may be unlocked.",
    "Founder separately starts archive-wide processing, if and when desired.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("12. Review statement", level=1)
p = doc.add_paragraph()
r = p.add_run("Current state: ")
r.bold = True
p.add_run("Draft definition prepared from verified PRD v0.2 for founder review. BUILD NOT AUTHORIZED. No application code, migrations, processing jobs, recognition runs, derived-data remediation, or runtime-data changes are authorized by this document.")

# Preserve the source header/footer construction but update visible identity.
for section in doc.sections:
    for paragraph in section.footer.paragraphs:
        if "MBPRD-P2-I13" in paragraph.text:
            paragraph.text = "MBBS-P2-I13  |  Definition Draft v0.2  |  Based on PRD v0.2"
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

doc.core_properties.title = "P2-I13 Increment Definition — Draft v0.2"
doc.core_properties.subject = "Founder-review definition derived from MBPRD-P2-I13 v0.2; build not authorized"
doc.core_properties.comments = "Draft for founder review. No build authorization."
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
normalize_docx_container(OUTPUT)
print(OUTPUT)
